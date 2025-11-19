from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

class DVRGenerator:
    """Classe per la generazione automatica del DVR"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_document_style()
    
    def setup_document_style(self):
        """Configura gli stili di base"""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    
    def add_header(self):
        """Aggiunge l'intestazione del documento"""
        # Titolo
        title = self.doc.add_paragraph()
        title_run = title.add_run("DOCUMENTO DI VALUTAZIONE DEI RISCHI")
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sottotitolo
        subtitle = self.doc.add_paragraph()
        subtitle_run = subtitle.add_run("(DVR)")
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(f"Data di compilazione: {datetime.now().strftime('%d/%m/%Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Spazio
    
    def add_company_section(self, company_data):
        """Aggiunge le informazioni aziendali"""
        self.doc.add_heading("1. INFORMAZIONI AZIENDALI", level=1)
        
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        company_info = [
            ("Ragione Sociale", company_data.get('name', 'N/A')),
            ("Indirizzo", company_data.get('address', 'N/A')),
            ("Settore Industriale", company_data.get('industry', 'N/A')),
            ("Numero Dipendenti", company_data.get('employees', 'N/A')),
            ("Responsabile Sicurezza", company_data.get('safety_manager', 'N/A'))
        ]
        
        for idx, (label, value) in enumerate(company_info):
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[1].text = str(value)
            table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_paragraph()
    
    def add_risk_matrix(self, risks):
        """Aggiunge la matrice di valutazione dei rischi"""
        self.doc.add_heading("2. MATRICE DI VALUTAZIONE DEI RISCHI", level=1)
        
        table = self.doc.add_table(rows=len(risks) + 1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Intestazioni
        headers = ['Rischio', 'Probabilità (P)', 'Gravità (G)', 'Esposizione (E)', 'Valore (P×G×E)', 'Livello']
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
            table.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        
        # Dati rischi
        for row_idx, risk in enumerate(risks, 1):
            cells = table.rows[row_idx].cells
            cells[0].text = risk['name']
            cells[1].text = str(risk['probability'])
            cells[2].text = str(risk['severity'])
            cells[3].text = str(risk['exposure'])
            
            risk_value = risk['probability'] * risk['severity'] * risk['exposure']
            cells[4].text = str(risk_value)
            
            # Classificazione del rischio
            if risk_value < 15:
                level = "BASSO"
            elif risk_value < 30:
                level = "MEDIO"
            else:
                level = "ALTO"
            cells[5].text = level
        
        self.doc.add_paragraph()
    
    def add_mitigation_section(self, mitigations):
        """Aggiunge le misure di mitigazione"""
        self.doc.add_heading("3. MISURE DI MITIGAZIONE", level=1)
        
        self.doc.add_heading("Misure Tecniche:", level=2)
        for measure in mitigations.get('technical', []):
            self.doc.add_paragraph(measure, style='List Bullet')
        
        self.doc.add_heading("Misure Organizzative:", level=2)
        for measure in mitigations.get('organizational', []):
            self.doc.add_paragraph(measure, style='List Bullet')
        
        self.doc.add_heading("Misure Procedurali:", level=2)
        for measure in mitigations.get('procedural', []):
            self.doc.add_paragraph(measure, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def add_compliance_section(self):
        """Aggiunge la sezione di conformità"""
        self.doc.add_heading("4. CONFORMITÀ NORMATIVA", level=1)
        
        norms = [
            "D.Lgs. 81/2008 - Testo Unico sulla Sicurezza del Lavoro",
            "GDPR (Regolamento UE 2016/679) - Protezione dei Dati",
            "Direttiva NIS2 - Sicurezza delle Reti e dei Sistemi Informativi",
            "D.Lgs. 231/2001 - Responsabilità Amministrativa degli Enti",
            "ISO/IEC 27001 - Gestione della Sicurezza delle Informazioni",
            "ISO 45001 - Gestione della Sicurezza e della Salute sul Lavoro"
        ]
        
        for norm in norms:
            self.doc.add_paragraph(norm, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def add_signature_section(self):
        """Aggiunge la sezione firme"""
        self.doc.add_heading("5. APPROVAZIONI E FIRME", level=1)
        
        table = self.doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        
        signatures = [
            "Datore di Lavoro",
            "Responsabile SPP (RSPP)",
            "Medico Competente",
            "Rappresentante dei Lavoratori (RLS)"
        ]
        
        for idx, sig in enumerate(signatures):
            table.rows[idx].cells[0].text = sig
            table.rows[idx].cells[1].text = "Firma: _______________"
            table.rows[idx].cells[2].text = "Data: _______________"
    
    def save(self, filepath):
        """Salva il documento"""
        self.doc.save(filepath)
        return os.path.exists(filepath)
